import io
import pdfplumber
import docx

class ExtractionError(Exception):
    pass

# extraction of text from pdf
def extract_text_from_pdf(file_bytes:bytes) -> str:
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()

                if (page_text):
                    text_parts.append(page_text)

    except Exception as e:
        raise ExtractionError(f"unable to read your pdf file: {e}")

    full_text = '\n'.join(text_parts).strip()

    if not full_text:
        raise ExtractionError('unable to extract text from your resume, try putting pdf and docx based resume')

    return full_text

 # extraction of text from docx
def extract_text_from_docx(file_bytes:bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ExtractionError(f'unable to read the document : {e}')

    parts = []

    #paragraphs
    for p in document.paragraphs:
        if (p.text.strip()):
            parts.append(p.text.strip())

    #tabular data
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if (cell.text.strip()):
                    parts.append(cell.text.strip())

    full_text = '\n'.join(parts).strip()

    if not full_text:
        raise ExtractionError('no extractable text from your docx file')

    return full_text

# here i will write code for checking file extension
def extract_resume_text(uploaded_file) -> str:

    if(uploaded_file is None):
        raise ExtractionError('No file provided yet')

    file_name = uploaded_file.name.lower()
    file_bytes = uploaded_file.getvalue()

    if not file_bytes:
        raise ExtractionError('uploaded file is empty')

    if (file_name.endswith('.pdf')):
        extract_text_from_pdf(file_bytes)
    elif(file_name.endswith('.docx')):
        extract_text_from_docx(file_bytes)
    else:
        raise ExtractionError('unsupported file type, please upload pdf or docx file type')
